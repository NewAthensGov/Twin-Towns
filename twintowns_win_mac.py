from typing import Union
import discord
from discord import option
from discord.ext import commands
from docx import Document
from datetime import datetime
from docx2pdf import convert
import asyncio
import os

intents = discord.Intents.default()
intents.message_content = True
bot = discord.Bot(intents=intents)

TOKEN = 'YOUR_DISCORD_BOT_TOKEN'
template_path = 'Twin Towns Certificate Template.docx'
certificates_dir = 'certificates'

# Ensure the certificates directory exists
os.makedirs(certificates_dir, exist_ok=True)


@bot.slash_command(name="twin")
@option("nation1", description="Enter the first nation")
@option("town1", description="Enter the first town")
@option("nation2", description="Enter the second nation")
@option("town2", description="Enter the second town")
@option("approver_user", description="Mention the approver user", type=discord.User)
async def twin(
    ctx: discord.ApplicationContext,
    nation1: str,
    town1: str,
    nation2: str,
    town2: str,
    approver_user: discord.User,
):
    response = await ctx.respond(f'{approver_user.mention}, please confirm the twin towns {town1}-{nation1} and {town2}-{nation2} with ✅ or decline with ❌.')
    message = await response.original_response()
    await message.add_reaction('✅')
    await message.add_reaction('❌')

    def check(reaction, user):
        return user == approver_user and str(reaction.emoji) in ['✅', '❌'] and reaction.message.id == message.id

    try:
        reaction, user = await bot.wait_for('reaction_add', timeout=60.0, check=check)
        if str(reaction.emoji) == '✅':
            await ctx.send_followup(f'Generating twin towns certificate for {town1}-{nation1} and {town2}-{nation2}...')
            output_docx_path = os.path.join(certificates_dir, f'{nation1}-{town1}_{nation2}-{town2}.docx')
            output_pdf_path = os.path.join(certificates_dir, f'{nation1}-{town1}_{nation2}-{town2}.pdf')
            submitter_nickname = ctx.author.nick if ctx.author.nick else ctx.author.name
            approver_member = ctx.guild.get_member(approver_user.id)
            approver_nickname = approver_member.nick if approver_member.nick else approver_member.name
            generate_certificate(town1, nation1, town2, nation2, submitter_nickname, approver_nickname, output_docx_path, output_pdf_path)
            await ctx.send_followup(file=discord.File(output_pdf_path))
        else:
            await ctx.send_followup('Twin towns request declined.')
    except asyncio.TimeoutError:
        await ctx.send_followup('Twin towns request expired.')

def replace_text_in_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        for run in paragraph.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

def replace_text_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, replacements)

def generate_certificate(town1, nation1, town2, nation2, submitter, approver, output_docx_path, output_pdf_path):
    document = Document(template_path)
    current_date = datetime.utcnow().strftime('%Y-%m-%d')
    replacements = {
        "{town1}": town1,
        "{nation1}": nation1,
        "{town2}": town2,
        "{nation2}": nation2,
        "{date}": current_date,
        "{submitter}": submitter,
        "{approver}": approver,
    }
    
    # Replace text in the main document
    replace_text_in_paragraphs(document.paragraphs, replacements)
    replace_text_in_tables(document.tables, replacements)
    
    # Replace text in headers and footers
    for section in document.sections:
        header = section.header
        footer = section.footer
        replace_text_in_paragraphs(header.paragraphs, replacements)
        replace_text_in_paragraphs(footer.paragraphs, replacements)
        replace_text_in_tables(header.tables, replacements)
        replace_text_in_tables(footer.tables, replacements)

    document.save(output_docx_path)

    # Convert DOCX to PDF using docx2pdf
    convert(output_docx_path, output_pdf_path)

bot.run(TOKEN)
